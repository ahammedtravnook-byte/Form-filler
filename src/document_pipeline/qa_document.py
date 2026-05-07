"""Q&A DOCX parsing via python-docx.

Extracts every question/answer pair from the document and returns them as a
plain dict keyed by the original question text. No fixed key mapping — the
caller decides how to interpret the answers.

Supported layouts (tried in order):
  1. Table-based  : 2+ column table, question in col 0, answer in col 1.
  2. Inline       : single paragraph "Key: Value" or "Key? Value"
                    — the most common Schengen questionnaire format.
  3. Labelled     : paragraphs prefixed with "Q:" / "Question:" / "1."
                    followed by a paragraph prefixed with "A:" / "Answer:".
  4. Adjacent     : alternating paragraphs — odd = question, even = answer
                    (only used when none of the above match).

Install: pip install python-docx
"""

from __future__ import annotations

import io
import re

from document_pipeline.models import QAData

_RE_Q_PREFIX = re.compile(r"^(?:Q(?:uestion)?\s*\d*|^\d+[\.\)])\s*[:\.\-]?\s*", re.IGNORECASE)
_RE_A_PREFIX = re.compile(r"^A(?:nswer)?\s*\d*\s*[:\.\-]?\s*", re.IGNORECASE)
# Inline "Key: Value" or "Key? Value". The key must contain at least one
# letter (so we don't catch things like "20.08.2026" or "P.O. Box: ...").
_RE_INLINE_KV = re.compile(r"^([^:?\n]{2,}?[A-Za-z][^:?\n]*)\s*[:?]\s*(.+)$")


def extract_qa_answers(docx_bytes: bytes) -> QAData:
    """Parse a .docx Q&A document and return all question/answer pairs.

    Keys in raw_answers are the verbatim question text (after stripping
    Q:/1. prefixes). Values are the corresponding answer text.
    """
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    doc = Document(io.BytesIO(docx_bytes))
    answers: dict[str, str] = {}

    # Strategy 1: table format
    for table in doc.tables:
        _parse_table(table, answers)

    # Strategies 2-4: paragraph-based
    if not answers:
        _parse_paragraphs(list(doc.paragraphs), answers)

    return QAData(raw_answers=answers)


def _parse_table(table: object, answers: dict[str, str]) -> None:
    from docx.table import Table  # type: ignore[import-untyped]

    assert isinstance(table, Table)
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 2:
            continue
        question_text = _strip_q_prefix(cells[0])
        answer_text = _strip_a_prefix(cells[1])
        if question_text and answer_text:
            answers[question_text] = answer_text


def _parse_paragraphs(paragraphs: list, answers: dict[str, str]) -> None:
    non_empty = [p.text.strip() for p in paragraphs if p.text.strip()]

    # Pass 1: inline "Key: Value" pairs — the most common Schengen layout.
    inline = _parse_inline_kv(non_empty)
    if inline:
        answers.update(inline)
        return

    # Pass 2: explicit Q:/A: labelled pairs.
    labelled: dict[str, str] = {}
    pending_q: str | None = None
    for text in non_empty:
        if _is_question(text):
            pending_q = _strip_q_prefix(text)
            continue
        if _is_answer(text) and pending_q:
            labelled[pending_q] = _strip_a_prefix(text)
            pending_q = None
            continue
        # A non-empty, non-labelled paragraph following a question is its answer.
        if pending_q:
            labelled[pending_q] = text
            pending_q = None

    if labelled:
        answers.update(labelled)
        return

    # Pass 3: alternating paragraphs (no prefixes at all).
    for i in range(0, len(non_empty) - 1, 2):
        answers[non_empty[i]] = non_empty[i + 1]


def _parse_inline_kv(lines: list[str]) -> dict[str, str]:
    """Parse "Key: Value" / "Key? Value" lines into a dict.

    Skips section headers (no separator) and blank values. We require at
    least 3 successful matches before returning anything, so a docx that
    happens to contain one stray "X: Y" sentence isn't classified as inline.
    """
    result: dict[str, str] = {}
    for line in lines:
        m = _RE_INLINE_KV.match(line)
        if not m:
            continue
        key = m.group(1).strip()
        value = m.group(2).strip()
        if key and value:
            result[key] = value
    return result if len(result) >= 3 else {}


def _is_question(text: str) -> bool:
    return bool(_RE_Q_PREFIX.match(text))


def _is_answer(text: str) -> bool:
    return bool(_RE_A_PREFIX.match(text))


def _strip_q_prefix(text: str) -> str:
    return _RE_Q_PREFIX.sub("", text).strip()


def _strip_a_prefix(text: str) -> str:
    return _RE_A_PREFIX.sub("", text).strip()
